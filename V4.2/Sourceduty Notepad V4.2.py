# Sourceduty Notepad V4.2
# Copyright (C) 2024, Sourceduty

import tkinter as tk
from tkinter.filedialog import askopenfilename, asksaveasfilename
from tkinter import simpledialog, messagebox
from tkinter import scrolledtext
import time
import os
import markdown
from docx import Document
import csv
import fitz  # PyMuPDF for PDF handling
import pandas as pd
import json
import html
import difflib
import random
import string

# Define templates including custom templates
TEMPLATES = {
    "Business": {
        "Meeting Notes": "Meeting Notes\nDate: \nAttendees: \nAgenda:\n- \n- \n\nDiscussion:\n- \n- \n\nAction Items:\n- \n- \n\nNext Meeting: ",
        "Project Plan": "Project Plan\nProject Name: \nStart Date: \nEnd Date: \nProject Manager: \n\nGoals:\n- \n- \n\nMilestones:\n- \n- \n\nResources:\n- \n- \n\nRisks:\n- \n- \n\nDeliverables:\n- \n- ",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "# Copyright (C) 20--, YourName\n\n[Add your content here]\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n",
        "Business Process": "Business Process\nTitle: \n\nObjective: \n\nSteps:\n1. Step 1 Description\n2. Step 2 Description\n3. Step 3 Description\n\nResources: \n- Resource 1\n- Resource 2\n\nOutcome: \n\n"
    },
    "Education": {
        "Lecture Notes": "Lecture Notes\nDate: \nInstructor: \nCourse: \n\nKey Points:\n- \n- \n\nSummary:\n\nQuestions:\n- \n- ",
        "Essay Outline": "Essay Outline\nTitle: \nThesis Statement: \n\nIntroduction:\n\nBody Paragraphs:\n1. \n2. \n3. \n\nConclusion: ",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "# Copyright (C) 20--, YourName\n\n[Add your content here]\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n"
    },
    "Creative Writing": {
        "Story Outline": "Story Outline\nTitle: \nGenre: \n\nCharacters:\n1. \n2. \n3. \n\nPlot Overview:\n\nChapter Breakdown:\n1. \n2. \n3. ",
        "Poem Template": "Poem Template\nTitle: \n\n[Write your poem here]\n\n\nReflection:\n\n",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "# Copyright (C) 20--, YourName\n\n[Add your content here]\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n"
    },
    "Custom Templates": {}
}

class TextEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Sourceduty Notepad V4.2")
        self.start_time = time.time()
        self.create_menu()
        self.create_widgets()
        self.set_default_description()
        self.update_footer()

    def create_menu(self):
        menu = tk.Menu(self.root)
        self.root.config(menu=menu)

        file_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New", command=self.new_file)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Open and Merge", command=self.open_and_merge)
        file_menu.add_command(label="Compare Files", command=self.compare_files)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_command(label="Save As...", command=self.save_as_file)
        file_menu.add_command(label="Export", command=self.export_file)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.exit)

        templates_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="Templates", menu=templates_menu)
        for industry, templates in TEMPLATES.items():
            industry_menu = tk.Menu(templates_menu, tearoff=0)
            templates_menu.add_cascade(label=industry, menu=industry_menu)
            for template_name in templates:
                industry_menu.add_command(label=template_name, command=lambda name=template_name, industry=industry: self.load_template(industry, name))

        custom_template_menu = tk.Menu(templates_menu, tearoff=0)
        templates_menu.add_cascade(label="Add Custom Template", menu=custom_template_menu)
        custom_template_menu.add_command(label="Add Template", command=self.add_custom_template)

        mode_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="Mode", menu=mode_menu)
        mode_menu.add_command(label="Dark Mode", command=self.dark_mode)        
        mode_menu.add_command(label="Light Mode", command=self.light_mode)

        control_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="Control", menu=control_menu)
        control_menu.add_command(label="About", command=self.show_options)
        control_menu.add_command(label="Text Statistics", command=self.show_statistics)
        control_menu.add_command(label="Generate Password", command=self.generate_password)
        control_menu.add_command(label="Sort/Filter and Export List", command=self.export_list)

    def create_widgets(self):
        self.text = tk.Text(self.root, wrap="word")
        self.text.pack(expand=1, fill="both")

        # Create a footer frame to contain the footer elements
        self.footer_frame = tk.Frame(self.root, bg="black")
        self.footer_frame.pack(side="bottom", fill="x")

        # Dynamic part of the footer (time)
        self.dynamic_footer = tk.Label(self.footer_frame, text="", anchor="w", bg="black", fg="white")
        self.dynamic_footer.pack(side="left", padx=10)

        # Timestamp button on the right side
        self.timestamp_button = tk.Button(self.footer_frame, text="Insert Timestamp", command=self.insert_timestamp, bg="black", fg="white")
        self.timestamp_button.pack(side="right", padx=10)

    def set_default_description(self):
        default_description = (
            "Sourceduty Notepad V4.2\n"
            "Copyright (C) 2024, Sourceduty\n\n"
            "Welcome to Sourceduty Notepad V4.2, your versatile text editor designed for various industry needs.\n\n"
            "Features:\n\n"
            "- Light and Dark Mode: Seamlessly switch between light and dark themes to match your working environment.\n"
            "- Template Management: Quickly load and manage templates tailored for business, education, and creative writing.\n"
            "- Custom Templates: Add up to 5 custom templates.\n"
            "- File Operations: Open, save, merge, and export files in multiple formats including TXT, DOCX, CSV, PDF, HTML, JSON, and more.\n"
            "- Text Statistics: Get character, word, and line counts for your text.\n"
            "- Password Generator: Generate secure passwords with ease.\n"
            "- Sort/Filter and Export List: Sort and filter your text and export it into a new .txt file list.\n"
            "- File Comparison: Compare two files to highlight differences.\n"
            "- Timestamp Insertion: Easily insert the current timestamp into your document with a single click.\n"
            "\n\nRepository: https://github.com/sourceduty/Notepad"
        )
        self.text.insert(tk.END, default_description)

    def update_footer(self):
        # Calculate elapsed time
        elapsed_time = time.time() - self.start_time
        minutes, seconds = divmod(int(elapsed_time), 60)

        # Create the dynamic footer text (time)
        time_str = f"Total Time: {minutes}m {seconds}s"

        # Update the dynamic footer only if the text has changed
        current_footer_text = self.dynamic_footer.cget("text")
        if current_footer_text != time_str:
            self.dynamic_footer.config(text=time_str)

        # Schedule the next update
        self.root.after(1000, self.update_footer)

    def insert_timestamp(self):
        timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
        self.text.insert(tk.END, f"\n{timestamp}")

    def new_file(self):
        self.text.delete(1.0, tk.END)
        self.start_time = time.time()

    def open_file(self):
        filepath = askopenfilename(filetypes=[("Text Files", "*.txt"), ("Markdown Files", "*.md"), ("Word Documents", "*.docx"), ("PDF Files", "*.pdf"), ("CSV Files", "*.csv"), ("HTML Files", "*.html"), ("JSON Files", "*.json"), ("All Files", "*.*")])
        if not filepath:
            return

        extension = os.path.splitext(filepath)[1].lower()
        if extension == ".txt":
            with open(filepath, "r") as input_file:
                text = input_file.read()
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, text)
        elif extension == ".md":
            with open(filepath, "r") as input_file:
                text = markdown.markdown(input_file.read())
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, text)
        elif extension == ".docx":
            doc = Document(filepath)
            full_text = [para.text for para in doc.paragraphs]
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, "\n".join(full_text))
        elif extension == ".pdf":
            doc = fitz.open(filepath)
            full_text = [page.get_text() for page in doc]
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, "\n".join(full_text))
        elif extension == ".csv":
            df = pd.read_csv(filepath)
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, df.to_string())
        elif extension == ".html":
            with open(filepath, "r") as input_file:
                text = html.unescape(input_file.read())
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, text)
        elif extension == ".json":
            with open(filepath, "r") as input_file:
                data = json.load(input_file)
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, json.dumps(data, indent=4))
        else:
            messagebox.showerror("Unsupported File Type", "The selected file type is not supported.")

    def open_and_merge(self):
        filepath = askopenfilename(filetypes=[("Text Files", "*.txt"), ("Markdown Files", "*.md"), ("Word Documents", "*.docx"), ("PDF Files", "*.pdf"), ("CSV Files", "*.csv"), ("HTML Files", "*.html"), ("JSON Files", "*.json"), ("All Files", "*.*")])
        if not filepath:
            return

        extension = os.path.splitext(filepath)[1].lower()
        existing_content = self.text.get(1.0, tk.END).strip()

        new_content = ""
        if extension == ".txt":
            with open(filepath, "r") as input_file:
                new_content = input_file.read()
        elif extension == ".md":
            with open(filepath, "r") as input_file:
                new_content = markdown.markdown(input_file.read())
        elif extension == ".docx":
            doc = Document(filepath)
            new_content = "\n".join([para.text for para in doc.paragraphs])
        elif extension == ".pdf":
            doc = fitz.open(filepath)
            new_content = "\n".join([page.get_text() for page in doc])
        elif extension == ".csv":
            df = pd.read_csv(filepath)
            new_content = df.to_string()
        elif extension == ".html":
            with open(filepath, "r") as input_file:
                new_content = html.unescape(input_file.read())
        elif extension == ".json":
            with open(filepath, "r") as input_file:
                data = json.load(input_file)
                new_content = json.dumps(data, indent=4)

        if new_content:
            self.text.delete(1.0, tk.END)
            merged_content = existing_content + "\n\n" + new_content
            self.text.insert(tk.END, merged_content)
            messagebox.showinfo("Merge Complete", "File merged successfully.")

    def compare_files(self):
        filepath1 = askopenfilename(title="Select the first file to compare")
        if not filepath1:
            return
        filepath2 = askopenfilename(title="Select the second file to compare")
        if not filepath2:
            return

        with open(filepath1, "r") as file1, open(filepath2, "r") as file2:
            content1 = file1.readlines()
            content2 = file2.readlines()

        diff = difflib.ndiff(content1, content2)
        diff_text = "\n".join(diff)

        comparison_window = tk.Toplevel(self.root)
        comparison_window.title("File Comparison")
        comparison_text = scrolledtext.ScrolledText(comparison_window, wrap="word")
        comparison_text.pack(expand=1, fill="both")
        comparison_text.insert(tk.END, diff_text)

    def save_file(self):
        filepath = asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("Markdown Files", "*.md"), ("Word Documents", "*.docx"), ("PDF Files", "*.pdf"), ("CSV Files", "*.csv"), ("HTML Files", "*.html"), ("JSON Files", "*.json"), ("All Files", "*.*")])
        if not filepath:
            return

        extension = os.path.splitext(filepath)[1].lower()
        content = self.text.get(1.0, tk.END).strip()

        if extension == ".txt":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".md":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".docx":
            doc = Document()
            doc.add_paragraph(content)
            doc.save(filepath)
        elif extension == ".pdf":
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), content)
            doc.save(filepath)
        elif extension == ".csv":
            with open(filepath, "w", newline="") as output_file:
                writer = csv.writer(output_file)
                for line in content.splitlines():
                    writer.writerow([line])
        elif extension == ".html":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".json":
            with open(filepath, "w") as output_file:
                json.dump(json.loads(content), output_file, indent=4)
        else:
            messagebox.showerror("Unsupported File Type", "The selected file type is not supported.")

    def save_as_file(self):
        self.save_file()

    def export_file(self):
        filepath = asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf"), ("Text Files", "*.txt"), ("Markdown Files", "*.md"), ("Word Documents", "*.docx"), ("CSV Files", "*.csv"), ("HTML Files", "*.html"), ("JSON Files", "*.json")])
        if not filepath:
            return

        content = self.text.get(1.0, tk.END).strip()
        extension = os.path.splitext(filepath)[1].lower()

        if extension == ".pdf":
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), content)
            doc.save(filepath)
        elif extension == ".txt":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".md":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".docx":
            doc = Document()
            doc.add_paragraph(content)
            doc.save(filepath)
        elif extension == ".csv":
            with open(filepath, "w", newline="") as output_file:
                writer = csv.writer(output_file)
                for line in content.splitlines():
                    writer.writerow([line])
        elif extension == ".html":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".json":
            with open(filepath, "w") as output_file:
                json.dump(json.loads(content), output_file, indent=4)
        else:
            messagebox.showerror("Unsupported File Type", "The selected file type is not supported.")

    def exit(self):
        self.root.quit()

    def load_template(self, industry, template_name):
        if industry in TEMPLATES and template_name in TEMPLATES[industry]:
            template = TEMPLATES[industry][template_name]
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, template)
        else:
            messagebox.showerror("Template Not Found", "The selected template was not found.")

    def add_custom_template(self):
        if len(TEMPLATES["Custom Templates"]) >= 5:
            messagebox.showerror("Limit Reached", "You can only add up to 5 custom templates.")
            return
        
        template_name = simpledialog.askstring("Template Name", "Enter the name for your custom template:")
        if template_name:
            content = self.text.get(1.0, tk.END).strip()
            TEMPLATES["Custom Templates"][template_name] = content
            messagebox.showinfo("Template Added", f"Custom template '{template_name}' added successfully.")

    def dark_mode(self):
        self.mode = "dark"
        self.update_mode()

    def light_mode(self):
        self.mode = "light"
        self.update_mode()

    def update_mode(self):
        if self.mode == "dark":
            self.root.config(bg="black")
            self.text.config(bg="black", fg="white")
            self.footer_frame.config(bg="black")
            self.dynamic_footer.config(bg="black", fg="white")
            self.timestamp_button.config(bg="black", fg="white")
        else:
            self.root.config(bg="white")
            self.text.config(bg="white", fg="black")
            self.footer_frame.config(bg="black")
            self.dynamic_footer.config(bg="black", fg="white")
            self.timestamp_button.config(bg="black", fg="white")

    def show_options(self):
        messagebox.showinfo("About", "Copyright (C) 2024, Sourceduty")

    def show_statistics(self):
        text = self.text.get(1.0, tk.END)
        num_characters = len(text)
        num_words = len(text.split())
        num_lines = text.count('\n') + 1
        messagebox.showinfo("Text Statistics", f"Characters: {num_characters}\nWords: {num_words}\nLines: {num_lines}")

    def generate_password(self):
        length = 12
        password = ''.join(random.choices(string.ascii_letters + string.digits + string.punctuation, k=length))
        self.text.delete(1.0, tk.END)
        self.text.insert(tk.END, f"Generated Password: {password}")

    def export_list(self):
        content = self.text.get(1.0, tk.END).splitlines()
        sorted_content = sorted(content)
        filtered_content = [line for line in sorted_content if line.strip()]  # Simple filter to remove empty lines

        filepath = asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
        if not filepath:
            return

        with open(filepath, "w") as output_file:
            output_file.write("\n".join(filtered_content))

        messagebox.showinfo("Export Complete", "Sorted and filtered list has been exported.")

if __name__ == "__main__":
    root = tk.Tk()
    app = TextEditor(root)
    root.mainloop()
