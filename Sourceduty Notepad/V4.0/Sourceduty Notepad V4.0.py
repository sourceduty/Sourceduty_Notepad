# Sourceduty Notepad V4.0
# Copyright (C) 2024, Sourceduty 

import tkinter as tk
from tkinter.filedialog import askopenfilename, asksaveasfilename
from tkinter import messagebox
import time
import os
import markdown
from docx import Document
import csv
import fitz 
import pandas as pd
import json
import html
import random
import string

TEMPLATES = {
    "Business": {
        "Meeting Notes": "Meeting Notes\nDate: \nAttendees: \nAgenda:\n- \n- \n\nDiscussion:\n- \n- \n\nAction Items:\n- \n- \n\nNext Meeting: ",
        "Project Plan": "Project Plan\nProject Name: \nStart Date: \nEnd Date: \nProject Manager: \n\nGoals:\n- \n- \n\nMilestones:\n- \n- \n\nResources:\n- \n- \n\nRisks:\n- \n- \n\nDeliverables:\n- \n- ",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "Copyright\nTitle: \nAuthor: \nYear: \n\nCopyright Notice:\n\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n"
    },
    "Education": {
        "Lecture Notes": "Lecture Notes\nDate: \nInstructor: \nCourse: \n\nKey Points:\n- \n- \n\nSummary:\n\nQuestions:\n- \n- ",
        "Essay Outline": "Essay Outline\nTitle: \nThesis Statement: \n\nIntroduction:\n\nBody Paragraphs:\n1. \n2. \n3. \n\nConclusion: ",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "Copyright\nTitle: \nAuthor: \nYear: \n\nCopyright Notice:\n\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n"
    },
    "Creative Writing": {
        "Story Outline": "Story Outline\nTitle: \nGenre: \n\nCharacters:\n1. \n2. \n3. \n\nPlot Overview:\n\nChapter Breakdown:\n1. \n2. \n3. ",
        "Poem Template": "Poem Template\nTitle: \n\n[Write your poem here]\n\n\nReflection:\n\n",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "Copyright\nTitle: \nAuthor: \nYear: \n\nCopyright Notice:\n\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n"
    }
}

class TextEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Sourceduty Notepad V4.0")
        self.start_time = time.time()
        self.create_menu()
        self.create_widgets()
        self.mode = "light"
        self.timer_log_enabled = True
        self.set_default_description()
        self.update_footer()

    def create_menu(self):
        menu = tk.Menu(self.root)
        self.root.config(menu=menu)

        file_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New", command=self.new_file)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_command(label="Save As...", command=self.save_as_file)
        file_menu.add_command(label="Export", command=self.export_file)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.exit)

        templates_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="Templates", menu=templates_menu)
        for industry, templates in TEMPLATES.items():
            industry_menu = tk.Menu(templates_menu, tearoff=0)
            templates_menu.add_cascade(label=industry, menu=industry_menu)
            for template_name in templates:
                industry_menu.add_command(label=template_name, command=lambda name=template_name, industry=industry: self.load_template(industry, name))

        mode_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="Mode", menu=mode_menu)
        mode_menu.add_command(label="Dark Mode", command=self.dark_mode)        
        mode_menu.add_command(label="Light Mode", command=self.light_mode)

        control_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="Control", menu=control_menu)
        control_menu.add_command(label="About", command=self.show_options)
        control_menu.add_command(label="Text Statistics", command=self.show_statistics)
        control_menu.add_command(label="Generate Password", command=self.generate_password)

    def create_widgets(self):
        self.text = tk.Text(self.root, wrap="word")
        self.text.pack(expand=1, fill="both")

        self.footer = tk.Label(self.root, text="", anchor="w", bg="lightgrey")
        self.footer.pack(side="bottom", fill="x")

    def set_default_description(self):
        default_description = (
            "Sourceduty Notepad V4.0\n"
            "Copyright (C) 2024, Sourceduty\n\n"
            "Features:\n"
            "- Light and Dark Mode: Seamlessly switch between light and dark themes to match your working environment.\n"
            "- Template Management: Quickly load and manage templates tailored for business, education, and creative writing.\n"
            "- File Operations: Open, save, and export files in multiple formats including TXT, DOCX, CSV, PDF, and more.\n"
            "- Text Statistics: Get character, word, and line counts for your text.\n"
            "- Password Generator: Generate secure passwords with ease.\n"
            "- Timer Log: Optionally track your time spent in each session.\n"
            "\n\nRepository: https://github.com/sourceduty/Notepad"
        )
        self.text.insert(tk.END, default_description)

    def new_file(self):
        self.text.delete(1.0, tk.END)
        self.start_time = time.time()

    def open_file(self):
        filepath = askopenfilename(filetypes=[("Text Files", "*.txt"), ("Markdown Files", "*.md"), ("Word Documents", "*.docx"), ("PDF Files", "*.pdf"), ("CSV Files", "*.csv"), ("HTML Files", "*.html"), ("JSON Files", "*.json"), ("All Files", "*.*")])
        if not filepath:
            return

        extension = os.path.splitext(filepath)[1].lower()
        if extension == ".txt":
            with open(filepath, "r") as input_file:
                text = input_file.read()
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, text)
        elif extension == ".md":
            with open(filepath, "r") as input_file:
                text = markdown.markdown(input_file.read())
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, text)
        elif extension == ".docx":
            doc = Document(filepath)
            full_text = [para.text for para in doc.paragraphs]
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, "\n".join(full_text))
        elif extension == ".pdf":
            doc = fitz.open(filepath)
            full_text = [page.get_text() for page in doc]
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, "\n".join(full_text))
        elif extension == ".csv":
            df = pd.read_csv(filepath)
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, df.to_string())
        elif extension == ".html":
            with open(filepath, "r") as input_file:
                text = html.unescape(input_file.read())
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, text)
        elif extension == ".json":
            with open(filepath, "r") as input_file:
                data = json.load(input_file)
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, json.dumps(data, indent=4))
        else:
            messagebox.showerror("Unsupported File Type", "The selected file type is not supported.")

    def save_file(self):
        filepath = asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("Markdown Files", "*.md"), ("Word Documents", "*.docx"), ("PDF Files", "*.pdf"), ("CSV Files", "*.csv"), ("HTML Files", "*.html"), ("JSON Files", "*.json"), ("All Files", "*.*")])
        if not filepath:
            return

        extension = os.path.splitext(filepath)[1].lower()
        content = self.text.get(1.0, tk.END).strip()

        if extension == ".txt":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".md":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".docx":
            doc = Document()
            doc.add_paragraph(content)
            doc.save(filepath)
        elif extension == ".pdf":
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), content)
            doc.save(filepath)
        elif extension == ".csv":
            with open(filepath, "w", newline="") as output_file:
                writer = csv.writer(output_file)
                for line in content.splitlines():
                    writer.writerow([line])
        elif extension == ".html":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".json":
            with open(filepath, "w") as output_file:
                json.dump(json.loads(content), output_file, indent=4)
        else:
            messagebox.showerror("Unsupported File Type", "The selected file type is not supported.")

    def save_as_file(self):
        self.save_file()

    def export_file(self):
        filepath = asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf"), ("Text Files", "*.txt"), ("Markdown Files", "*.md"), ("Word Documents", "*.docx"), ("CSV Files", "*.csv"), ("HTML Files", "*.html"), ("JSON Files", "*.json")])
        if not filepath:
            return

        content = self.text.get(1.0, tk.END).strip()
        extension = os.path.splitext(filepath)[1].lower()

        if extension == ".pdf":
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), content)
            doc.save(filepath)
        elif extension == ".txt":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".md":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".docx":
            doc = Document()
            doc.add_paragraph(content)
            doc.save(filepath)
        elif extension == ".csv":
            with open(filepath, "w", newline="") as output_file:
                writer = csv.writer(output_file)
                for line in content.splitlines():
                    writer.writerow([line])
        elif extension == ".html":
            with open(filepath, "w") as output_file:
                output_file.write(content)
        elif extension == ".json":
            with open(filepath, "w") as output_file:
                json.dump(json.loads(content), output_file, indent=4)
        else:
            messagebox.showerror("Unsupported File Type", "The selected file type is not supported.")

    def exit(self):
        self.root.quit()

    def load_template(self, industry, template_name):
        if industry in TEMPLATES and template_name in TEMPLATES[industry]:
            template = TEMPLATES[industry][template_name]
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, template)
        else:
            messagebox.showerror("Template Not Found", "The selected template was not found.")

    def dark_mode(self):
        self.mode = "dark"
        self.update_mode()

    def light_mode(self):
        self.mode = "light"
        self.update_mode()

    def update_mode(self):
        if self.mode == "dark":
            self.root.config(bg="black")
            self.text.config(bg="black", fg="white")
            self.footer.config(bg="black", fg="white")
        else:
            self.root.config(bg="white")
            self.text.config(bg="white", fg="black")
            self.footer.config(bg="white", fg="black")

    def show_options(self):
        messagebox.showinfo("About", "Copyright (C) 2024, Sourceduty")

    def show_statistics(self):
        text = self.text.get(1.0, tk.END)
        num_characters = len(text)
        num_words = len(text.split())
        num_lines = text.count('\n') + 1
        messagebox.showinfo("Text Statistics", f"Characters: {num_characters}\nWords: {num_words}\nLines: {num_lines}")

    def generate_password(self):
        length = 12
        password = ''.join(random.choices(string.ascii_letters + string.digits + string.punctuation, k=length))
        self.text.delete(1.0, tk.END)
        self.text.insert(tk.END, f"Generated Password: {password}")

    def update_footer(self):
        elapsed_time = time.time() - self.start_time
        minutes, seconds = divmod(int(elapsed_time), 60)
        time_str = f"Total Time: {minutes}m {seconds}s | Line: {self.text.index(tk.INSERT).split('.')[0]}"
        self.footer.config(text=time_str)
        self.root.after(1000, self.update_footer)

root = tk.Tk()
app = TextEditor(root)
root.mainloop()
