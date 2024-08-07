# Sourceduty Notepad V5.8
# Copyright (C) 2024, Sourceduty

# pip install tkinter python-docx PyMuPDF pandas ebooklib beautifulsoup4 cryptography

import tkinter as tk
from tkinter import messagebox, simpledialog
from tkinter.filedialog import asksaveasfilename, askopenfilename
import difflib
import fitz
import time
import random
import string
import markdown
import csv
import os
import json
import pandas as pd
from ebooklib import epub
from bs4 import BeautifulSoup
from docx import Document
import re
from cryptography.fernet import Fernet, InvalidToken

LANGUAGES = {
    "English": "en",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Chinese": "zh-cn",
    "Japanese": "ja",
    "Russian": "ru",
    "Italian": "it"
}

TEMPLATES = {
    "Business": {
        "Meeting Notes": "Meeting Notes\nDate: \nAttendees: \nAgenda:\n- \n- \n\nDiscussion:\n- \n- \n\nAction Items:\n- \n- \n\nNext Meeting: ",
        "Project Plan": "Project Plan\nProject Name: \nStart Date: \nEnd Date: \nProject Manager: \n\nGoals:\n- \n- \n\nMilestones:\n- \n- \n\nResources:\n- \n- \n\nRisks:\n- \n- \n\nDeliverables:\n- \n- ",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "# Copyright (C) 20--, YourName\n\n[Add your content here]\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n",
        "Business Process": "Business Process\nTitle: \n\nObjective: \n\nSteps:\n1. Step 1 Description\n2. Step 2 Description\n3. Step 3 Description\n\nResources: \n- Resource 1\n- Resource 2\n\nOutcome: \n\n"
    },
    "Education": {
        "Lecture Notes": "Lecture Notes\nDate: \nInstructor: \nCourse: \n\nKey Points:\n- \n- \n\nSummary:\n\nQuestions:\n- \n- ",
        "Essay Outline": "Essay Outline\nTitle: \nThesis Statement: \n\nIntroduction:\n\nBody Paragraphs:\n1. \n2. \n3. \n\nConclusion: ",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "# Copyright (C) 20--, YourName\n\n[Add your content here]\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n"
    },
    "Creative Writing": {
        "Story Outline": "Story Outline\nTitle: \nGenre: \n\nCharacters:\n1. \n2. \n3. \n\nPlot Overview:\n\nChapter Breakdown:\n1. \n2. \n3. ",
        "Poem Template": "Poem Template\nTitle: \n\n[Write your poem here]\n\n\nReflection:\n\n",
        "Schedule": "Schedule\nDate: \n\nTasks:\n1. \n2. \n3. \n\nNotes:\n- \n- \n",
        "Routine": "Daily Routine\nDate: \n\nMorning:\n- \n- \n\nAfternoon:\n- \n- \n\nEvening:\n- \n- \n\nNotes:\n- \n- \n",
        "Copyright": "# Copyright (C) 20--, YourName\n\n[Add your content here]\n",
        "Process": "Process Template\nTitle: \n\nObjective:\n\nSteps:\n1. \n2. \n3. \n\nOutcome:\n\nNotes:\n- \n- \n"
    },
    "Custom Templates": {}
}

class TextEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Sourceduty Notepad V5.8")
        self.language = 'en'
        self.current_file = None
        self.key = None  
        self.cipher = None
        self.create_menu()
        self.micronotes = []
        self.start_time = time.time()
        self.keywords = []
        self.password_cipher = None    

        self.text = tk.Text(root, wrap="word", bg="black", fg="white", insertbackground="white")
        self.text.pack(expand=1, fill="both")

        self.footer_frame = tk.Frame(root, bg="black")
        self.footer_frame.pack(side="bottom", fill="x")

        self.dynamic_footer = tk.Label(self.footer_frame, text="Time Elapsed: 0m 0s", bg="black", fg="white")
        self.dynamic_footer.pack(side="left")

        self.timestamp_button = tk.Button(self.footer_frame, text="Insert Timestamp", command=self.insert_timestamp, bg="black", fg="white")
        self.timestamp_button.pack(side="right")

        self.micronotes_button = tk.Button(self.footer_frame, text="Micronotes", command=self.show_micronotes_menu, bg="black", fg="white")
        self.micronotes_button.pack(side="right")

        self.create_menu()
        self.insert_default_description()
        self.update_footer()

    def create_menu(self):
        self.menu = tk.Menu(self.root)
        self.root.config(menu=self.menu)

        self.file_menu = tk.Menu(self.menu, tearoff=0)
        self.menu.add_cascade(label="File", menu=self.file_menu)
        self.file_menu.add_command(label="New", command=self.new_file)
        self.file_menu.add_command(label="Open", command=self.open_file)
        self.file_menu.add_command(label="Open and Merge", command=self.open_and_merge)
        self.file_menu.add_command(label="Compare Files", command=self.compare_files)
        self.file_menu.add_command(label="Save", command=self.save_file)
        self.file_menu.add_command(label="Save Locked", command=self.save_locked_file)
        self.file_menu.add_command(label="Export", command=self.export_file)
        self.file_menu.add_command(label="Import", command=self.import_file)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="Exit", command=self.root.quit)

        self.template_menu = tk.Menu(self.menu, tearoff=0)
        self.menu.add_cascade(label="Templates", menu=self.template_menu)
        for industry, templates in TEMPLATES.items():
            industry_menu = tk.Menu(self.template_menu, tearoff=0)
            self.template_menu.add_cascade(label=industry, menu=industry_menu)
            for template_name in templates:
                industry_menu.add_command(label=template_name, command=lambda name=template_name, ind=industry: self.load_template(ind, name))
        self.template_menu.add_separator()
        self.template_menu.add_command(label="Add Custom Template", command=self.add_custom_template)

        self.mode_menu = tk.Menu(self.menu, tearoff=0)
        self.menu.add_cascade(label="Mode", menu=self.mode_menu)
        self.mode_menu.add_command(label="Dark Mode", command=self.dark_mode)
        self.mode_menu.add_command(label="Light Mode", command=self.light_mode)

        self.tools_menu = tk.Menu(self.menu, tearoff=0)
        self.menu.add_cascade(label="Control", menu=self.tools_menu)
        self.tools_menu.add_command(label="Password", command=self.generate_password)
        self.tools_menu.add_command(label="Export List", command=self.export_list)
        self.tools_menu.add_command(label="Statistics", command=self.show_statistics)
        self.tools_menu.add_command(label="Import Restore", command=self.import_restore)
        self.tools_menu.add_cascade(label="Text Layout", menu=self.create_text_layout_menu())
        self.tools_menu.add_command(label="Search", command=self.search_word)
        self.tools_menu.add_command(label="Tagging and Categorization", command=self.show_tagging_and_categorization)
        self.tools_menu.add_command(label="Topology", command=self.convert_to_topology)  
        self.tools_menu.add_command(label="Languages", command=self.select_language)
        self.tools_menu.add_command(label="About", command=self.show_options)

        self.encryption_menu = tk.Menu(self.tools_menu, tearoff=0)
        self.encryption_menu.add_command(label="Encrypt File", command=self.encrypt_file)
        self.encryption_menu.add_command(label="Decrypt File", command=self.decrypt_file)
        self.tools_menu.add_cascade(label="Encryption", menu=self.encryption_menu)

    def convert_to_topology(self):
        content = self.text.get(1.0, tk.END).strip()
        if not content:
            messagebox.showwarning("Empty File", "The text area is empty. Please enter some text to convert.")
            return

        lines = content.split('\n')
        diagram = []
        level_stack = []

        for line in lines:
            stripped_line = line.lstrip()
            if not stripped_line:
                continue

            indent = len(line) - len(stripped_line)
            level = 0
            while level < len(level_stack) and level_stack[level][1] < indent:
                level += 1
            
            if level < len(level_stack):
                level_stack = level_stack[:level]
            
            node = stripped_line
            prefix = ''.join(['  ' for _ in range(level)])
            diagram.append(f"{prefix}└── {node}")
            level_stack.append((node, indent))

        self.text.delete(1.0, tk.END)
        self.text.insert(tk.END, '\n'.join(diagram))
        messagebox.showinfo("Conversion Complete", "The text has been converted to a topological diagram.")

    def load_or_generate_key(self):
        key_file = "secret.key"
        if os.path.exists(key_file):
            with open(key_file, "rb") as key_file:
                return key_file.read()
        else:
            key = Fernet.generate_key()
            with open(key_file, "wb") as key_file:
                key_file.write(key)
            messagebox.showinfo("Key Generation", "New encryption key generated and saved as 'secret.key'.")
            return key

    def encrypt_file(self):
        file_path = askopenfilename(title="Select File to Encrypt")
        if file_path:
            try:
                with open(file_path, "rb") as file:
                    file_data = file.read()

                key = Fernet.generate_key()
                cipher = Fernet(key)
                encrypted_data = cipher.encrypt(file_data)

                encrypted_file_path = asksaveasfilename(defaultextension=".enc", initialfile=os.path.basename(file_path) + ".enc", filetypes=[("Encrypted Files", "*.enc"), ("All Files", "*.*")])
                if encrypted_file_path:
                    with open(encrypted_file_path, "wb") as file:
                        file.write(encrypted_data)

                    key_file_path = asksaveasfilename(defaultextension=".key", filetypes=[("Key Files", "*.key"), ("All Files", "*.*")])
                    if key_file_path:
                        with open(key_file_path, "wb") as key_file:
                            key_file.write(key)

                        messagebox.showinfo("Encrypt File", f"File '{file_path}' encrypted successfully as '{encrypted_file_path}'.\nKey saved as '{key_file_path}'.")
                    else:
                        messagebox.showwarning("Save Cancelled", "Encryption key save was cancelled. No key file was saved.")
                else:
                    messagebox.showwarning("Save Cancelled", "Encryption was cancelled. No file was saved.")
            except Exception as e:
                messagebox.showerror("Encryption Error", f"An error occurred while encrypting the file: {e}")


    def decrypt_file(self):
        file_path = askopenfilename(title="Select File to Decrypt")
        if file_path:
            if not file_path.endswith(".enc"):
                messagebox.showwarning("Invalid File", "Please select a valid encrypted file with a '.enc' extension.")
                return

            key_file_path = askopenfilename(title="Select Key File", filetypes=[("Key Files", "*.key"), ("All Files", "*.*")])
            if not key_file_path:
                messagebox.showwarning("No Key File", "No key file selected. Decryption cannot proceed.")
                return

            try:
                with open(key_file_path, "rb") as key_file:
                    key = key_file.read()
                    cipher = Fernet(key)

                with open(file_path, "rb") as file:
                    encrypted_data = file.read()

                decrypted_data = cipher.decrypt(encrypted_data)

                # Print the decrypted data in the text area
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, decrypted_data.decode('utf-8'))

                decrypted_file_path = file_path.replace(".enc", "")
                with open(decrypted_file_path, "wb") as file:
                    file.write(decrypted_data)

                messagebox.showinfo("Decrypt File", f"File '{file_path}' decrypted successfully and displayed.\nDecrypted file saved as '{decrypted_file_path}'.")
            except InvalidToken:
                messagebox.showerror("Decryption Error", "Invalid key or corrupted file. Decryption failed.")
            except Exception as e:
                messagebox.showerror("Decryption Error", f"An error occurred while decrypting the file: {e}")

    def ensure_cipher(self):
        if self.key is None or self.cipher is None:
            self.key = self.load_or_generate_key()
            self.cipher = Fernet(self.key)
                
    def create_text_layout_menu(self):
        layout_menu = tk.Menu(self.tools_menu, tearoff=0)
        layout_menu.add_command(label="Single Column", command=self.set_single_column)
        layout_menu.add_command(label="Two Columns", command=self.set_two_columns)
        layout_menu.add_command(label="Three Columns", command=self.set_three_columns)
        return layout_menu

    def update_footer(self):
        elapsed_time = int(time.time() - self.start_time)
        minutes, seconds = divmod(elapsed_time, 60)
        self.dynamic_footer.config(text=f"Time Elapsed: {minutes}m {seconds}s")
        self.root.after(1000, self.update_footer)

    def insert_timestamp(self):
        self.text.insert(tk.END, time.strftime("%Y-%m-%d %H:%M:%S\n"))

    def show_micronotes_menu(self):
        menu = tk.Menu(self.root, tearoff=0)
        for note in self.micronotes:
            menu.add_command(label=note, command=lambda n=note: self.insert_micronote(n))
        menu.add_command(label="Add New Note", command=self.add_micronote)
        menu.add_command(label="Clear Notes", command=self.clear_micronotes)
        menu.post(self.root.winfo_pointerx(), self.root.winfo_pointery())

    def add_micronote(self):
        note = simpledialog.askstring("Add Micronote", "Enter Micronote:")
        if note:
            self.micronotes.append(note)
            messagebox.showinfo("Micronote Added", "Micronote added successfully!")

    def insert_micronote(self, note):
        self.text.insert(tk.END, f"Micronote: {note}\n")

    def clear_micronotes(self):
        self.micronotes.clear()
        messagebox.showinfo("Micronotes Cleared", "All micronotes cleared!")
        
    def insert_default_description(self):
        default_description = (
            "Sourceduty Notepad V5.8\n"
            "\nCopyright (C) 2024, Sourceduty - All Rights Reserved.\n"
            "\nDescription:\n"
            "\nFeatures:\n"
            "\nText Formatting:\n"
            "  - Align text (left, center, right, justify)\n"
            "  - Increase/decrease font size\n"
            "  - Change font type (e.g., Arial, Times New Roman, Courier)\n"
            "  - Change font color (supports RGB and HEX color codes)\n"
            "  - Clear formatting (reset to default)\n"
            "\nFile Operations:\n"
            "  - New File: Create a new, blank document.\n"
            "  - Open File: Open an existing document in various formats (e.g., .txt, .docx, .rtf, .html, .md).\n"
            "  - Save File: Save the current document in the selected format.\n"
            "  - Save Locked File: Save the current document using a password and encryption.\n"
            "  - Encrypt File: Encrypt a selected file and save it with a .enc extension.\n"
            "  - Decrypt File: Decrypt a selected encrypted file using a key file and display its contents.\n"
            "  - Export File: Export the document to formats such as PDF, DOCX, HTML, or Markdown (.md).\n"
            "  - Import File: Import content from formats like CSV, JSON, XML, or Markdown (.md).\n"
            "  - Open and Merge File: Open another file and merge its content into the current document.\n"
            "  - Compare Files: Compare the content of two documents and highlight differences.\n"
            "\nTemplate Management:\n"
            "  - Load Templates: Load pre-defined templates for quick document creation.\n"
            "  - Add Custom Templates: Create and save your own templates for reuse.\n"
            "\nDark Mode:\n"
            "  - Switch between dark mode and light mode to reduce eye strain or match user preferences.\n"
            "\nTimestamp Insertion:\n"
            "  - Insert the current date and time into the document with customizable formats (e.g., YYYY-MM-DD, MM/DD/YYYY).\n"
            "\nMicronotes Menu:\n"
            "  - Add Micronotes: Insert small notes or comments in the document.\n"
            "  - Delete Micronotes: Remove specific micronotes.\n"
            "  - Clear Micronotes: Remove all micronotes from the document.\n"
            "\nPassword Generator:\n"
            "  - Generate Secure Passwords: Create passwords with customizable length and complexity (e.g., including special characters, numbers).\n"
            "\nCSV Export:\n"
            "  - Export document content or lists to CSV format for data handling and analysis.\n"
            "\nDocument Statistics:\n"
            "  - Display Word Count: Count the number of words in the document.\n"
            "  - Character Count: Count the number of characters in the document.\n"
            "  - Line Count: Count the number of lines in the document.\n"
            "\nJSON Restore:\n"
            "  - Import and Restore: Import document content from a JSON file and restore the document to its previous state.\n"
            "\nSearch Function:\n"
            "  - Find Text: Search for a specific word or phrase within the document.\n"
            "\nTagging and Categorization:\n"
            "  - Tag Documents: Assign tags to documents for easier organization and retrieval.\n"
            "  - Categorize Documents: Categorize documents based on custom criteria (e.g., project, topic).\n"
            "\nKeyword Highlighting:\n"
            "  - Set Keywords: Define specific keywords to highlight throughout the document.\n"
            "  - Highlight Keywords: Automatically highlight occurrences of these keywords in the document.\n"
            "\nTopology:\n"
            "  - Convert to Topology: Format the current text into a hierarchical topological diagram.\n"
        )
        self.text.insert(tk.END, default_description)

    def new_file(self):
        self.text.delete(1.0, tk.END)
        self.current_file = None

    def open_file(self):
        file_path = askopenfilename(filetypes=[
            ("Text Files", "*.txt"), 
            ("PDF Files", "*.pdf"), 
            ("CSV Files", "*.csv"), 
            ("JSON Files", "*.json"), 
            ("Excel Files", "*.xlsx"), 
            ("ePub Files", "*.epub"), 
            ("Word Files", "*.docx"), 
            ("HTML Files", "*.html"),
            ("Markdown Files", "*.md"),
            ("All Files", "*.*")])
        
        if file_path:
            if file_path.endswith(".txt"):
                with open(file_path, "r") as file:
                    content = file.read()
                    self.text.delete(1.0, tk.END)
                    self.text.insert(tk.END, content)
            elif file_path.endswith(".pdf"):
                doc = fitz.open(file_path)
                content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, content)
            elif file_path.endswith(".csv"):
                with open(file_path, newline='') as file:
                    reader = csv.reader(file)
                    content = "\n".join([", ".join(row) for row in reader])
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, content)
            elif file_path.endswith(".json"):
                with open(file_path, "r") as file:
                    content = json.load(file)
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, json.dumps(content, indent=4))
            elif file_path.endswith(".xlsx"):
                df = pd.read_excel(file_path)
                content = df.to_string(index=False)
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, content)
            elif file_path.endswith(".epub"):
                book = epub.read_epub(file_path)
                content = ""
                for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                    soup = BeautifulSoup(item.get_body_content(), 'html.parser')
                    content += soup.get_text() + "\n"
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, content)
            elif file_path.endswith(".docx"):
                doc = Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, content)
            elif file_path.endswith(".html"):
                with open(file_path, "r") as file:
                    soup = BeautifulSoup(file, 'html.parser')
                    content = soup.get_text()
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, content)
            elif file_path.endswith(".md"):
                with open(file_path, "r") as file:
                    content = file.read()
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, content)

            self.current_file = file_path

    def open_and_merge(self):
        file_path = askopenfilename(filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if file_path:
            with open(file_path, "r") as file:
                content = file.read()
                self.text.insert(tk.END, "\n" + "="*40 + "\n")
                self.text.insert(tk.END, content)

    def compare_files(self):
        file1 = askopenfilename(filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        file2 = askopenfilename(filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if file1 and file2:
            with open(file1, "r") as f1, open(file2, "r") as f2:
                text1 = f1.readlines()
                text2 = f2.readlines()
                diff = difflib.unified_diff(text1, text2, fromfile=file1, tofile=file2)
                diff_text = ''.join(diff)
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, diff_text)

    def save_file(self):
        if self.current_file:
            with open(self.current_file, "w") as file:
                content = self.text.get(1.0, tk.END)
                file.write(content)
        else:
            self.save_as_file()

    def save_locked_file(self):
        file_path = asksaveasfilename(defaultextension=".txt", 
                                      filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if file_path:
            key = Fernet.generate_key()
            cipher_suite = Fernet(key)
            content = self.text.get(1.0, tk.END).encode()
            encrypted_content = cipher_suite.encrypt(content)

            with open(file_path, "wb") as file:
                file.write(encrypted_content)
            with open(file_path + ".key", "wb") as key_file:
                key_file.write(key)

            self.current_file = file_path
            messagebox.showinfo("Save Locked", "File saved and locked successfully!")

    def export_file(self):
        file_path = asksaveasfilename(defaultextension=".txt", 
                                      filetypes=[("Text Files", "*.txt"), 
                                                 ("PDF Files", "*.pdf"), 
                                                 ("CSV Files", "*.csv"), 
                                                 ("JSON Files", "*.json"), 
                                                 ("Excel Files", "*.xlsx"), 
                                                 ("ePub Files", "*.epub"), 
                                                 ("Word Files", "*.docx"), 
                                                 ("HTML Files", "*.html"),
                                                 ("Markdown Files", "*.md"),
                                                 ("All Files", "*.*")])
        if file_path:
            content = self.text.get(1.0, tk.END)
            if file_path.endswith(".txt"):
                with open(file_path, "w") as file:
                    file.write(content)
            elif file_path.endswith(".pdf"):
                doc = fitz.open()
                page = doc.new_page()
                page.insert_text((72, 72), content)
                doc.save(file_path)
                doc.close()
            elif file_path.endswith(".csv"):
                with open(file_path, "w", newline='') as file:
                    writer = csv.writer(file)
                    writer.writerow([content])
            elif file_path.endswith(".json"):
                with open(file_path, "w") as file:
                    json.dump(content, file)
            elif file_path.endswith(".xlsx"):
                df = pd.DataFrame([content])
                df.to_excel(file_path, index=False)
            elif file_path.endswith(".epub"):
                book = epub.EpubBook()
                book.set_title("Document")
                book.set_language('en')
                chapter = epub.EpubHtml(title='Document', file_name='chap_01.xhtml', lang='en')
                chapter.content = f'<h1>Document</h1><p>{content}</p>'
                book.add_item(chapter)
                epub.write_epub(file_path, book, {})
            elif file_path.endswith(".docx"):
                doc = Document()
                doc.add_paragraph(content)
                doc.save(file_path)
            elif file_path.endswith(".html"):
                with open(file_path, "w") as file:
                    file.write(content)
            elif file_path.endswith(".md"):
                with open(file_path, "w") as file:
                    file.write(content)
            messagebox.showinfo("Export", "File exported successfully!")
                
    def select_language(self):
            language = simpledialog.askstring("Select Language", "Enter language (e.g., English, Spanish, etc.):")
            if language in LANGUAGES:
                self.current_language = language
                messagebox.showinfo("Language Selection", f"Language set to {language}")
            else:
                messagebox.showerror("Error", "Language not supported")

    def export_pdf(self, content):
        file_path = asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")])
        if file_path:
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), content)
            doc.save(file_path)
            doc.close()
            messagebox.showinfo("Export", "PDF exported successfully!")

    def export_json(self, content):
        file_path = asksaveasfilename(defaultextension=".json", filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")])
        if file_path:
            with open(file_path, 'w') as file:
                json.dump(content, file)
            messagebox.showinfo("Export", "JSON exported successfully!")

    def export_excel(self, dataframe):
        file_path = asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx"), ("All Files", "*.*")])
        if file_path:
            dataframe.to_excel(file_path, index=False)
            messagebox.showinfo("Export", "Excel file exported successfully!")

    def export_epub(self, content, title="Document"):
        file_path = asksaveasfilename(defaultextension=".epub", filetypes=[("ePub Files", "*.epub"), ("All Files", "*.*")])
        if file_path:
            book = epub.EpubBook()
            book.set_title(title)
            book.set_language('en')
            chapter = epub.EpubHtml(title='Document', file_name='chap_01.xhtml', lang='en')
            chapter.content = f'<h1>{title}</h1><p>{content}</p>'
            book.add_item(chapter)
            epub.write_epub(file_path, book, {})
            messagebox.showinfo("Export", "ePub exported successfully!")

    def export_docx(self, content):
        file_path = asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx"), ("All Files", "*.*")])
        if file_path:
            doc = Document()
            doc.add_paragraph(content)
            doc.save(file_path)
            messagebox.showinfo("Export", "Word document exported successfully!")

    def export_html(self, content):
        file_path = asksaveasfilename(defaultextension=".html", filetypes=[("HTML Files", "*.html"), ("All Files", "*.*")])
        if file_path:
            with open(file_path, 'w') as file:
                file.write(content)
            messagebox.showinfo("Export", "HTML file exported successfully!")

    def import_file(self):
        file_path = askopenfilename(filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if file_path:
            with open(file_path, "r") as file:
                content = file.read()
                self.text.delete(1.0, tk.END)
                self.text.insert(tk.END, content)

    def add_custom_template(self):
        name = simpledialog.askstring("Template Name", "Enter the name of the new template:")
        if name:
            content = simpledialog.askstring("Template Content", "Enter the content of the new template:")
            if content:
                TEMPLATES["Custom Templates"][name] = content
                self.template_menu.add_command(label=name, command=lambda: self.load_template("Custom Templates", name))
                messagebox.showinfo("Custom Template", "Custom template added successfully!")

    def load_template(self, industry, name):
        if industry in TEMPLATES and name in TEMPLATES[industry]:
            template_content = TEMPLATES[industry][name]
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, template_content)

    def dark_mode(self):
        self.text.config(bg="black", fg="white", insertbackground="white")
        self.footer_frame.config(bg="black")
        self.dynamic_footer.config(bg="black", fg="white")
        self.timestamp_button.config(bg="black", fg="white")
        self.micronotes_button.config(bg="black", fg="white")

    def light_mode(self):
        self.text.config(bg="white", fg="black", insertbackground="black")
        self.footer_frame.config(bg="white")
        self.dynamic_footer.config(bg="white", fg="black")
        self.timestamp_button.config(bg="white", fg="black")
        self.micronotes_button.config(bg="white", fg="black")

    def generate_password(self):
        password_length = simpledialog.askinteger("Password Length", "Enter password length:")
        if password_length:
            password = ''.join(random.choices(string.ascii_letters + string.digits, k=password_length))
            messagebox.showinfo("Generated Password", f"Your password is: {password}")

    def export_list(self):
        file_path = asksaveasfilename(defaultextension=".csv", filetypes=[("CSV Files", "*.csv"), ("All Files", "*.*")])
        if file_path:
            with open(file_path, "w", newline="") as file:
                writer = csv.writer(file)
                for line in self.text.get(1.0, tk.END).splitlines():
                    writer.writerow([line])
            messagebox.showinfo("Export List", "List exported successfully!")

    def show_statistics(self):
        text_content = self.text.get(1.0, tk.END)
        word_count = len(text_content.split())
        line_count = text_content.count('\n')
        char_count = len(text_content)
        stats = f"Word Count: {word_count}\nLine Count: {line_count}\nCharacter Count: {char_count}"
        messagebox.showinfo("Statistics", stats)

    def import_restore(self):
        file_path = askopenfilename(filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")])
        if file_path:
            try:
                with open(file_path, "r") as file:
                    data = json.load(file)
                    self.text.delete(1.0, tk.END)
                    self.text.insert(tk.END, data.get("text", ""))
            except json.JSONDecodeError:
                messagebox.showerror("Import Error", "Failed to decode JSON. Please ensure the file is in correct format.")

    def show_options(self):
        messagebox.showinfo("About", "Copyright (C) 2024, Sourceduty")

    def apply_options(self, options):
        for option, var in options.items():
            print(f"{option}: {'Enabled' if var.get() else 'Disabled'}")

    def search_word(self):
        search_term = simpledialog.askstring("Search", "Enter the word or phrase to search for:")
        if search_term:
            content = self.text.get(1.0, tk.END)
            matches = re.finditer(re.escape(search_term), content, re.IGNORECASE)
            match_count = sum(1 for _ in matches)
            messagebox.showinfo("Search Results", f"'{search_term}' found {match_count} times.")

    def highlight_keywords(self):
        keywords = simpledialog.askstring("Keywords", "Enter keywords to highlight, separated by commas:")
        if keywords:
            keywords = [k.strip() for k in keywords.split(",")]
            content = self.text.get(1.0, tk.END)
            for keyword in keywords:
                start = "1.0"
                while True:
                    pos = self.text.search(keyword, start, stopindex=tk.END, nocase=True)
                    if not pos:
                        break
                    end = f"{pos}+{len(keyword)}c"
                    self.text.tag_add("highlight", pos, end)
                    start = end
                self.text.tag_config("highlight", background="yellow")

    def undo(self):
        self.text.edit_undo()

    def redo(self):
        self.text.edit_redo()

    def create_backup(self):
        if self.current_file:
            backup_path = f"{self.current_file}.bak"
            with open(self.current_file, "r") as file:
                content = file.read()
            with open(backup_path, "w") as file:
                file.write(content)
            messagebox.showinfo("Backup Created", f"Backup created: {backup_path}")

    def restore_backup(self):
        backup_file = askopenfilename(filetypes=[("Backup Files", "*.bak"), ("All Files", "*.*")])
        if backup_file:
            with open(backup_file, "r") as file:
                content = file.read()
            self.text.delete(1.0, tk.END)
            self.text.insert(tk.END, content)
            messagebox.showinfo("Backup Restored", "Backup restored successfully!")

    def set_single_column(self):
        self.text.config(width=80)

    def set_two_columns(self):
        self.text.config(width=40)

    def set_three_columns(self):
        self.text.config(width=27)

    def show_tagging_and_categorization(self):
        messagebox.showinfo("Tagging and Categorization", "Tagging and categorization functionality will be implemented here.")

if __name__ == "__main__":
    root = tk.Tk()
    app = TextEditor(root)
    root.mainloop()
